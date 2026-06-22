#!/usr/bin/env python3
"""Extract N32G033 User Manual PDF to Word document."""
import PyPDF2
import os
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout.reconfigure(encoding='utf-8')

pdf_path = r'C:\Users\jgldzbaeghuwqibvb\Desktop\GMJS\CN_UM_N32G033_Series_User_Manual_V1.0.0.pdf'
output_dir = r'e:\N32\N32G033\docs'
os.makedirs(output_dir, exist_ok=True)

reader = PyPDF2.PdfReader(pdf_path)
total = len(reader.pages)
print(f'Total pages: {total}')

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10)

for i in range(total):
    page = reader.pages[i]
    text = page.extract_text()

    # Add page marker
    heading = doc.add_heading(f'Page {i+1}', level=2)

    if text and text.strip():
        # Split by paragraphs and add
        for para_text in text.split('\n'):
            para_text = para_text.strip()
            if para_text:
                p = doc.add_paragraph(para_text)
                p.style.font.size = Pt(9)
    else:
        doc.add_paragraph('[此页无可提取文本 / No extractable text]')

    if (i+1) % 50 == 0:
        print(f'Processed {i+1}/{total} pages...')

output_path = os.path.join(output_dir, 'N32G033_User_Manual.docx')
doc.save(output_path)
print(f'Done! Output: {output_path}')
print(f'File size: {os.path.getsize(output_path) / 1024 / 1024:.1f} MB')
